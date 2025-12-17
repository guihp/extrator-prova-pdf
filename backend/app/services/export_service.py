"""
Serviço para exportar questões em PDF e Word
"""
import os
from typing import List, Dict, Optional
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportService:
    """Serviço para exportar questões em diferentes formatos"""
    
    def __init__(self):
        pass
    
    def export_to_pdf(self, questoes: List[Dict], imagens: List[Dict] = None, 
                     prova_nome: str = "Prova") -> BytesIO:
        """
        Exporta questões para PDF
        
        Args:
            questoes: Lista de questões
            imagens: Lista de imagens (opcional)
            prova_nome: Nome da prova
            
        Returns:
            BytesIO com o PDF gerado
        """
        buffer = BytesIO()
        
        # Criar documento PDF
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilo para título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#000000',
            spaceAfter=30,
            alignment=TA_LEFT
        )
        
        # Estilo para número da questão
        questao_num_style = ParagraphStyle(
            'QuestaoNum',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#1a1a1a',
            spaceAfter=12,
            spaceBefore=20,
            alignment=TA_LEFT
        )
        
        # Estilo para texto da questão
        questao_text_style = ParagraphStyle(
            'QuestaoText',
            parent=styles['Normal'],
            fontSize=11,
            textColor='#333333',
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            leading=14
        )
        
        # Conteúdo do PDF
        story = []
        
        # Título
        story.append(Paragraph(f"<b>{prova_nome}</b>", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Adicionar questões
        for questao in questoes:
            numero = questao.get("numero", questao.get("ordem", 0))
            texto = questao.get("texto", "")
            
            # Número da questão
            story.append(Paragraph(
                f"<b>Questão {numero}</b>",
                questao_num_style
            ))
            
            # Texto da questão
            # Escapar HTML e quebrar linhas
            texto_html = texto.replace('\n', '<br/>')
            texto_html = texto_html.replace('&', '&amp;')
            texto_html = texto_html.replace('<', '&lt;')
            texto_html = texto_html.replace('>', '&gt;')
            
            story.append(Paragraph(texto_html, questao_text_style))
            story.append(Spacer(1, 0.3*inch))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
    
    def export_to_word(self, questoes: List[Dict], imagens: List[Dict] = None,
                      prova_nome: str = "Prova") -> BytesIO:
        """
        Exporta questões para Word (DOCX)
        
        Args:
            questoes: Lista de questões
            imagens: Lista de imagens (opcional)
            prova_nome: Nome da prova
            
        Returns:
            BytesIO com o DOCX gerado
        """
        # Criar documento Word
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Título
        title = doc.add_heading(prova_nome, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Adicionar questões
        for questao in questoes:
            numero = questao.get("numero", questao.get("ordem", 0))
            texto = questao.get("texto", "")
            
            # Número da questão
            heading = doc.add_heading(f'Questão {numero}', level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Texto da questão
            para = doc.add_paragraph(texto)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Espaço entre questões
            doc.add_paragraph()
        
        # Salvar em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def export_questao_individual_pdf(self, questao: Dict, imagens: List[Dict] = None) -> BytesIO:
        """
        Exporta uma questão individual para PDF
        
        Args:
            questao: Dicionário com dados da questão
            imagens: Lista de imagens relacionadas (opcional)
            
        Returns:
            BytesIO com o PDF gerado
        """
        return self.export_to_pdf([questao], imagens, f"Questão {questao.get('numero', 'N/A')}")
    
    def export_questao_individual_word(self, questao: Dict, imagens: List[Dict] = None) -> BytesIO:
        """
        Exporta uma questão individual para Word
        
        Args:
            questao: Dicionário com dados da questão
            imagens: Lista de imagens relacionadas (opcional)
            
        Returns:
            BytesIO com o DOCX gerado
        """
        return self.export_to_word([questao], imagens, f"Questão {questao.get('numero', 'N/A')}")


export_service = ExportService()



